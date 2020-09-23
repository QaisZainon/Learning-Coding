#! python3
# Generates a Word document with custom invitation from guest.txt
# Add styles to a blank Word file
# Invitation per page, use add_break()
import docx
from docx.shared import Pt
from docx.shared import Inches

doc = docx.Document('Invitation.docx')
txt = open('guests.txt', 'r')
guest = txt.readlines()
inv_txt = [
    'It would be a pleasure to have the company of',
    'at 11010 Memory Lane on the Evening of',
    'April 1st',
    'at 7 o\'clock'
    ]

i = 5
for guest_num in range(len(guest)):
    doc.add_paragraph(inv_txt[0], style = 'Style1').alignment = 1
    doc.add_paragraph(guest[guest_num], style = 'Style2').alignment = 1
    doc.add_paragraph(inv_txt[1], style = 'Style1').alignment = 1
    doc.add_paragraph(inv_txt[2], style = 'Style2').alignment = 1
    doc.add_paragraph(inv_txt[3], style = 'Style1').alignment = 1
    doc.paragraphs[i].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    i += 5

doc.save('Invitation.docx')